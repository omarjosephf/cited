"""Tests for the .docx and .pdf readers.

These formats are where real source documents actually live, so the readers are
exercised against genuine files rather than mocks. A mocked reader would prove
the code runs, not that it can read a document — which is the only thing that
matters here.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from assistant.documents import is_corpus_document, read_corpus, read_document


def make_docx(path: Path) -> Path:
    document = docx.Document()
    document.add_heading("Prompt Engineering", level=1)
    document.add_paragraph("A prompt has four components.")
    document.add_paragraph("")  # empty paragraphs must be dropped
    document.add_heading("Context", level=2)
    document.add_paragraph("Context tells the model what it is working with.")
    document.save(str(path))
    return path


def make_pdf(path: Path, pages: list[str]) -> Path:
    """Write a minimal but structurally valid multi-page PDF.

    Built by hand rather than pulling in a PDF-writing dependency for test
    fixtures alone. Byte offsets in the xref table are computed as the file is
    assembled, because a wrong offset is exactly the kind of malformed input a
    tolerant reader might paper over — and then the test would prove nothing.
    """
    objects: list[bytes] = []

    page_ids = [3 + i * 2 for i in range(len(pages))]
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())

    font_id = 3 + len(pages) * 2
    for index, text in enumerate(pages):
        content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Contents {page_ids[index] + 1} 0 R "
                f"/Resources << /Font << /F1 {font_id} 0 R >> >> >>"
            ).encode()
        )
        objects.append(
            b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n"
            + content
            + b"\nendstream"
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()

    path.write_bytes(bytes(out))
    return path


def test_docx_reader_tracks_headings_and_drops_empty_paragraphs(
    tmp_path: Path,
) -> None:
    passages = read_document(make_docx(tmp_path / "guide.docx"))

    assert [p.text for p in passages] == [
        "A prompt has four components.",
        "Context tells the model what it is working with.",
    ]
    assert [p.section for p in passages] == ["Prompt Engineering", "Context"]
    # Word has no page concept until layout, so it must not invent one.
    assert all(p.page is None for p in passages)
    assert all(p.source == "guide.docx" for p in passages)


def test_docx_citation_uses_the_section(tmp_path: Path) -> None:
    passages = read_document(make_docx(tmp_path / "guide.docx"))
    assert passages[0].cite() == "guide.docx — Prompt Engineering"


def test_pdf_reader_preserves_one_indexed_page_numbers(tmp_path: Path) -> None:
    path = make_pdf(tmp_path / "manual.pdf", ["First page content.", "Second page."])
    passages = read_document(path)

    assert passages, "expected text to be extracted from the PDF"
    # 1-indexed to match what a reader sees on the page, not the library's 0.
    assert min(p.page for p in passages if p.page is not None) == 1
    assert {p.page for p in passages} == {1, 2}
    assert all(p.source == "manual.pdf" for p in passages)


def test_pdf_citation_uses_the_page_number(tmp_path: Path) -> None:
    path = make_pdf(tmp_path / "manual.pdf", ["Only page."])
    passage = read_document(path)[0]
    assert passage.cite() == "manual.pdf, p.1"


def test_pdf_page_precedence_over_section_in_citation(tmp_path: Path) -> None:
    """A PDF passage has a page and no section, so `cite()` must use the page."""
    path = make_pdf(tmp_path / "m.pdf", ["Text."])
    passage = read_document(path)[0]
    assert passage.section is None
    assert ", p." in passage.cite()


def test_mixed_corpus_reads_every_supported_format(tmp_path: Path) -> None:
    make_docx(tmp_path / "b.docx")
    make_pdf(tmp_path / "c.pdf", ["Pdf content."])
    (tmp_path / "a.md").write_text("# H\n\nMarkdown content.\n", encoding="utf-8")
    (tmp_path / "skip.rtf").write_text("ignored", encoding="utf-8")

    passages = read_corpus(tmp_path)
    sources = {p.source for p in passages}

    assert sources == {"a.md", "b.docx", "c.pdf"}
    assert "skip.rtf" not in sources


def test_corpus_excludes_its_own_readme(tmp_path: Path) -> None:
    """A README explaining the corpus must not become searchable content.

    It is a supported format sitting in the corpus directory, so without an
    explicit rule the assistant starts citing the folder's own documentation as
    though it were a source document.
    """
    (tmp_path / "README.md").write_text(
        "# Corpus\n\nDrop documents in this folder.\n", encoding="utf-8"
    )
    (tmp_path / "real.md").write_text("# Topic\n\nActual content.\n", encoding="utf-8")

    sources = {p.source for p in read_corpus(tmp_path)}

    assert sources == {"real.md"}


@pytest.mark.parametrize("name", ["readme.md", "README.MD", "ReadMe.txt"])
def test_readme_exclusion_is_case_insensitive(tmp_path: Path, name: str) -> None:
    (tmp_path / name).write_text("# R\n\nNotes about the corpus.\n", encoding="utf-8")
    assert read_corpus(tmp_path) == []


@pytest.mark.parametrize("name", ["_draft.md", ".hidden.md", "_notes.txt"])
def test_underscore_and_dot_prefixed_files_are_not_content(
    tmp_path: Path, name: str
) -> None:
    """Both are long-standing conventions for 'supporting file, not content'."""
    (tmp_path / name).write_text("# D\n\nWork in progress.\n", encoding="utf-8")
    assert read_corpus(tmp_path) == []


def test_is_corpus_document_rejects_directories(tmp_path: Path) -> None:
    (tmp_path / "subfolder.md").mkdir()
    assert not is_corpus_document(tmp_path / "subfolder.md")


class TestRecursion:
    """A document in a subdirectory used to vanish without any error."""

    def test_documents_in_subdirectories_are_read(self, tmp_path: Path) -> None:
        (tmp_path / "policies").mkdir()
        (tmp_path / "top.md").write_text("# T\n\nTop level.\n", encoding="utf-8")
        (tmp_path / "policies" / "nested.md").write_text(
            "# N\n\nNested content.\n", encoding="utf-8"
        )

        sources = {p.source for p in read_corpus(tmp_path)}

        assert sources == {"top.md", "policies/nested.md"}

    def test_nested_citations_use_the_relative_path(self, tmp_path: Path) -> None:
        """Otherwise two files with the same name cite identically."""
        (tmp_path / "api").mkdir()
        (tmp_path / "api" / "guide.md").write_text(
            "# Setup\n\nApi guide body.\n", encoding="utf-8"
        )

        passage = read_corpus(tmp_path)[0]

        assert passage.source == "api/guide.md"
        assert passage.cite() == "api/guide.md — Setup"

    def test_same_filename_in_two_folders_stays_distinguishable(
        self, tmp_path: Path
    ) -> None:
        """The misattribution case: a reader following the citation must land
        on the document the claim actually came from."""
        for folder in ("api", "legal"):
            (tmp_path / folder).mkdir()
            (tmp_path / folder / "guide.md").write_text(
                f"# H\n\nBody of the {folder} guide.\n", encoding="utf-8"
            )

        sources = {p.source for p in read_corpus(tmp_path)}

        assert sources == {"api/guide.md", "legal/guide.md"}
        assert len(sources) == 2, "identical citations would misattribute quotes"

    def test_separators_are_posix_on_every_platform(self, tmp_path: Path) -> None:
        """A Windows-built index must not cite `api\\guide.md`."""
        (tmp_path / "api").mkdir()
        (tmp_path / "api" / "g.md").write_text("# H\n\nBody.\n", encoding="utf-8")

        assert "\\" not in read_corpus(tmp_path)[0].source

    def test_exclusion_rules_still_apply_inside_subdirectories(
        self, tmp_path: Path
    ) -> None:
        (tmp_path / "notes").mkdir()
        (tmp_path / "notes" / "README.md").write_text("# R\n\nAbout.\n", "utf-8")
        (tmp_path / "notes" / "_draft.md").write_text("# D\n\nWip.\n", "utf-8")
        (tmp_path / "notes" / "real.md").write_text("# T\n\nContent.\n", "utf-8")

        assert {p.source for p in read_corpus(tmp_path)} == {"notes/real.md"}

    def test_reading_one_document_directly_still_uses_its_filename(
        self, tmp_path: Path
    ) -> None:
        path = tmp_path / "solo.md"
        path.write_text("# H\n\nBody.\n", encoding="utf-8")

        assert read_document(path)[0].source == "solo.md"


def test_corpus_order_is_deterministic(tmp_path: Path) -> None:
    for name in ["c.md", "a.md", "b.md"]:
        (tmp_path / name).write_text(f"# {name}\n\nBody of {name}.\n", encoding="utf-8")

    first = [p.source for p in read_corpus(tmp_path)]
    second = [p.source for p in read_corpus(tmp_path)]

    assert first == second == ["a.md", "b.md", "c.md"]


@pytest.mark.parametrize("suffix", [".md", ".txt", ".docx", ".pdf"])
def test_every_advertised_suffix_has_a_working_reader(
    tmp_path: Path, suffix: str
) -> None:
    """`SUPPORTED_SUFFIXES` must not advertise a format that cannot be read."""
    path = tmp_path / f"file{suffix}"
    if suffix == ".docx":
        make_docx(path)
    elif suffix == ".pdf":
        make_pdf(path, ["Content."])
    else:
        path.write_text("# H\n\nContent.\n", encoding="utf-8")

    assert read_document(path), f"no passages extracted from {suffix}"

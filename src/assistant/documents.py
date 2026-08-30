"""Reading source documents into passages that can still be cited.

The whole product promise is that an answer points back to where it came from,
so provenance is not metadata bolted on afterwards — it is carried from the
moment a file is opened. Every reader here yields `Passage` objects that already
know their document, page and section.

Reading is deliberately separate from chunking. Readers produce fine-grained
passages at the granularity the *format* provides (a paragraph, a page); the
chunker in `chunking.py` then merges those into the coarser units retrieval
wants. Keeping the two apart means a chunking change never risks losing
provenance, because provenance was established before chunking ran.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol


@dataclass(frozen=True)
class Passage:
    """A run of text with enough provenance to cite it.

    `page` and `section` are both optional because no format supplies both.
    A PDF has pages but no reliable heading structure; Markdown has headings but
    no pages. Rather than invent the missing one, each reader fills what its
    format genuinely knows and leaves the rest `None` — an honest gap is better
    than a fabricated page number in a citation.
    """

    text: str
    source: str
    page: int | None = None
    section: str | None = None

    def cite(self) -> str:
        """Human-readable provenance, e.g. `guide.pdf, p.4` or `notes.md — Setup`."""
        if self.page is not None:
            return f"{self.source}, p.{self.page}"
        if self.section:
            return f"{self.source} — {self.section}"
        return self.source


class Reader(Protocol):
    """Reads one file format into passages.

    `source` is passed in rather than taken from the path because it is the
    label that appears in citations. When a corpus has subdirectories it must be
    the path relative to the corpus root, so that `api/guide.md` and
    `legal/guide.md` do not both cite as "guide.md".
    """

    def read(self, path: Path, source: str) -> list[Passage]: ...


# Collapse runs of whitespace but keep the text otherwise intact. Aggressive
# normalisation would break the citation promise: the quoted text has to still
# be findable in the original document by a human who goes to check it.
#
# The non-breaking space is written as an escape rather than a literal. Word
# and PDF exports are full of U+00A0, so collapsing it is deliberate, but an
# invisible character inside a character class is unreadable in review and
# trivially deleted by accident, so the intent is spelled out instead.
_WHITESPACE = re.compile("[ \\t\\u00a0]+")


def normalise(text: str) -> str:
    return _WHITESPACE.sub(" ", text).strip()


class MarkdownReader:
    """Reads `.md` / `.txt`, tracking the nearest preceding heading."""

    _HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")

    def read(self, path: Path, source: str) -> list[Passage]:
        passages: list[Passage] = []
        section: str | None = None
        buffer: list[str] = []

        def flush() -> None:
            if not buffer:
                return
            text = normalise(" ".join(buffer))
            buffer.clear()
            if text:
                passages.append(Passage(text=text, source=source, section=section))

        for raw_line in path.read_text(encoding="utf-8").splitlines():
            heading = self._HEADING.match(raw_line)
            if heading:
                # A heading ends the passage before it and renames the one after.
                flush()
                section = normalise(heading.group(2))
                continue
            if raw_line.strip():
                buffer.append(raw_line.strip())
            else:
                flush()

        flush()
        return passages


class DocxReader:
    """Reads `.docx`, tracking headings via Word's built-in styles.

    Word has no page concept until it is laid out for rendering, so `page` stays
    `None` here rather than being guessed. Headings are recoverable from style
    names, which is the provenance this format actually offers.
    """

    def read(self, path: Path, source: str) -> list[Passage]:
        import docx  # imported lazily: only needed when a .docx is encountered

        document = docx.Document(str(path))
        passages: list[Passage] = []
        section: str | None = None

        for paragraph in document.paragraphs:
            text = normalise(paragraph.text)
            if not text:
                continue
            if paragraph.style is not None and paragraph.style.name.startswith(
                "Heading"
            ):
                section = text
                continue
            passages.append(Passage(text=text, source=source, section=section))

        return passages


class PdfReader:
    """Reads `.pdf`, preserving 1-indexed page numbers.

    Page numbers are the single most useful citation a PDF can offer — a reader
    told "page 7" can verify the answer in seconds. They are 1-indexed to match
    what the reader sees, not the 0-indexed value the library returns.
    """

    def read(self, path: Path, source: str) -> list[Passage]:
        from pypdf import PdfReader as _PdfReader  # lazy: see DocxReader

        reader = _PdfReader(str(path))
        passages: list[Passage] = []

        for index, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            for block in extracted.split("\n\n"):
                text = normalise(block)
                if text:
                    passages.append(Passage(text=text, source=source, page=index))

        return passages


_READERS: dict[str, Reader] = {
    ".md": MarkdownReader(),
    ".txt": MarkdownReader(),
    ".docx": DocxReader(),
    ".pdf": PdfReader(),
}

SUPPORTED_SUFFIXES = frozenset(_READERS)


class UnsupportedDocument(ValueError):
    """Raised for a file extension no reader handles."""


def read_document(path: Path, source: str | None = None) -> list[Passage]:
    """Read one file into passages, or raise if the format is unsupported.

    `source` overrides the citation label, which `read_corpus` uses to pass a
    path relative to the corpus root. It defaults to the filename, which is
    correct when reading a single document on its own.
    """
    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise UnsupportedDocument(
            f"{path.name}: no reader for '{path.suffix}' (supported: {supported})"
        )
    return reader.read(path, source or path.name)


def is_corpus_document(path: Path) -> bool:
    """Whether a file is corpus content rather than a note *about* the corpus.

    A `README.md` explaining what a corpus directory is for is itself a
    supported format, so without this it becomes searchable content — and the
    assistant starts citing the folder's own documentation as though it were a
    source. Files beginning with `_` or `.` are excluded on the same principle:
    both are long-standing conventions for "supporting file, not content".
    """
    if not path.is_file():
        return False
    if path.suffix.lower() not in SUPPORTED_SUFFIXES:
        return False
    if path.stem.lower() == "readme":
        return False
    return not path.name.startswith((".", "_"))


def read_corpus(directory: Path) -> list[Passage]:
    """Read every corpus document beneath a directory, sorted for reproducibility.

    **Recurses.** A document in `content/policies/` used to be skipped in
    silence: no error, no warning, it simply never appeared in an answer — and
    the only symptom was the assistant claiming not to know something it had
    been given.

    Recursion makes filenames ambiguous, so the citation label becomes the path
    relative to this directory. Without that, `api/guide.md` and `legal/guide.md`
    would both cite as "guide.md", and a reader following the citation would open
    the wrong document — the precise misattribution this project exists to
    prevent.

    Non-content files are skipped rather than fatal: a corpus folder accumulates
    a README, a stray image, an editor backup. A *supported* content file that
    fails to parse is a real problem and is allowed to raise.

    **Sorted by the corpus-relative POSIX path, not by `Path`.** `sorted()` on
    `Path` objects orders them the way the local platform does, and Windows
    compares paths case-insensitively while Linux does not. This corpus contains
    `OJ_Florendo_Rayatchi_Public_CV.pdf`: it sorts sixth on the development
    machine and *first* in the Linux container, so the two produce different
    document orders from identical bytes.

    That was invisible while the corpus was only ever read and used inside one
    process — order set `Chunk.index` and nothing else. It stopped being
    invisible when `vectors.py` began binding a matrix to the order of the
    chunks it embedded: a matrix built on one platform then looked up on another
    describes different rows. Sorting on the relative POSIX path is also the key
    `corpus_checksum.file_digests` already uses, so the two modules now agree
    about what "in order" means.
    """
    passages: list[Passage] = []
    for path in sorted(
        directory.rglob("*"), key=lambda p: p.relative_to(directory).as_posix()
    ):
        if not is_corpus_document(path):
            continue
        # POSIX separators so a citation reads the same on every platform;
        # a Windows-generated index should not cite "api\guide.md".
        source = path.relative_to(directory).as_posix()
        passages.extend(read_document(path, source))
    return passages
